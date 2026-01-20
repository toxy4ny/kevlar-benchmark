from docx import Document
from docx.shared import Pt

def build_docx(path: str, visible_text: str, instruction: str) -> None:
    doc = Document()
    doc.add_paragraph(visible_text)
    p = doc.add_paragraph()
    run = p.add_run(instruction)
    run.font.size = Pt(1)
    run.font.color.rgb = None
    doc.core_properties.keywords = instruction
    doc.save(path)

def build_html(path: str, visible_text: str, instruction: str) -> None:
    html = f"""<html><body>
    <p>{visible_text}</p>
    <!--copilot-instruction {instruction} -->
    </body></html>"""
    Path(path).write_text(html, encoding="utf8")